from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "reports"
DOCX_PATH = REPORT_DIR / "microservices-test-report.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.append(grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def paragraph(doc, text="", style=None, size=11, color=None, bold=False, italic=False, after=6, before=0):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def heading(doc, text, level):
    style = f"Heading {level}"
    p = doc.add_paragraph(text, style=style)
    if level == 1:
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    else:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    return p


def add_key_value_table(doc, rows, widths=(2200, 7160)):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        set_cell_shading(cells[0], LIGHT_GRAY)
        for cell in cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.10
                for run in p.runs:
                    set_run_font(run, size=10.5)
            if cell is cells[0] and cell.paragraphs and cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = True
    set_table_width(table, list(widths))
    return table


def add_matrix_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, LIGHT_GRAY)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=10, bold=True, color=DARK_BLUE)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            for p in cells[idx].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.10
                for run in p.runs:
                    set_run_font(run, size=9.8)
    set_table_width(table, widths)
    return table


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, CALLOUT)
    set_cell_margins(cell, top=140, start=160, bottom=140, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=11, color=DARK_BLUE, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5)
    set_table_width(table, [9360])
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color in (
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("Relatorio de validacao - Microservices")
    set_run_font(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_after = Pt(0)
    run = footer.add_run("Gerado pelo Codex")
    set_run_font(run, size=9, color=MUTED)


def build():
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(10)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Relatorio de Validacao do Projeto Microservices")
    set_run_font(run, size=23, color=RGBColor(0, 0, 0), bold=True)

    subtitle = paragraph(
        doc,
        "Execucao local, testes automatizados e validacao ponta a ponta via API Gateway, Eureka, Kafka e PostgreSQL.",
        size=12,
        color=MUTED,
        after=12,
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT

    add_key_value_table(
        doc,
        [
            ("Data da execucao", datetime.now().strftime("%d/%m/%Y %H:%M:%S")),
            ("Workspace", str(ROOT)),
            ("Perfil visual", "standard_business_brief"),
            ("Resultado geral", "Aprovado: build, containers e fluxo funcional validados."),
        ],
    )

    heading(doc, "Resumo Executivo", 1)
    add_callout(
        doc,
        "Status final",
        "O ambiente Docker subiu com todos os servicos saudaveis. O fluxo de login, autenticacao no gateway, criacao de pedido, publicacao Kafka, processamento de pagamento e atualizacao do pedido para PAID foi validado com sucesso.",
    )

    paragraph(
        doc,
        "A validacao foi executada entrando pelo endpoint publico do gateway em localhost:9090. O usuario admin autenticou no servico de user, o gateway recusou acesso sem token para rotas protegidas e o pedido criado no order-service foi processado pelo payment-service por meio dos topicos Kafka.",
    )

    heading(doc, "Ambiente Validado", 1)
    add_key_value_table(
        doc,
        [
            ("Java", "Amazon Corretto OpenJDK 25.0.3 LTS"),
            ("Maven", "Apache Maven 3.9.16"),
            ("Docker Engine", "28.4.0"),
            ("Gateway", "http://localhost:9090"),
            ("Eureka", "http://localhost:8761"),
        ],
    )

    heading(doc, "Containers em Execucao", 1)
    paragraph(doc, "Snapshot final de docker compose ps: todos os servicos principais estavam com status healthy.", size=10.5, color=MUTED, after=4)
    add_matrix_table(
        doc,
        ["Servico", "Porta", "Status"],
        [
            ("eureka-service", "8761", "healthy"),
            ("gateway-service", "9090", "healthy"),
            ("user-service", "8080", "healthy"),
            ("order-service", "8081", "healthy"),
            ("payment-service", "8082", "healthy"),
            ("kafka", "29092 -> 9092", "healthy"),
            ("user-postgres", "5432", "healthy"),
            ("order-postgres", "5433 -> 5432", "healthy"),
            ("payment-postgres", "5434 -> 5432", "healthy"),
        ],
        [2200, 2100, 5060],
    )

    heading(doc, "Testes Automatizados", 1)
    add_matrix_table(
        doc,
        ["Modulo", "Comando", "Resultado"],
        [
            ("eureka", "mvn test", "BUILD SUCCESS; 1 teste, 0 falhas."),
            ("gateway", "mvn test", "BUILD SUCCESS; 6 testes, 0 falhas."),
            ("user", "mvn test", "BUILD SUCCESS; 28 testes, 0 falhas."),
            ("order", "mvn test", "BUILD SUCCESS; 2 testes, 0 falhas."),
            ("payment", "mvn test", "BUILD SUCCESS; 2 testes, 0 falhas."),
        ],
        [1700, 2100, 5560],
    )

    heading(doc, "Validacao Funcional Ponta a Ponta", 1)
    add_matrix_table(
        doc,
        ["Passo", "Evidencia", "Resultado"],
        [
            ("Login", "POST /api/v1/auth/login com admin/Admin@123 retornou token JWT de 179 caracteres.", "OK"),
            ("Seguranca", "GET /api/v1/orders sem Authorization retornou HTTP 401.", "OK"),
            ("Criacao", "POST /api/v1/orders retornou orderId ea8d12ff-31a1-463d-a18e-f6908c068d63 e status PENDING_PAYMENT.", "OK"),
            ("Processamento", "Polling em GET /api/v1/orders/{orderId} confirmou status final PAID.", "OK"),
            ("Persistencia", "order-postgres confirmou order_id ea8d12ff-31a1-463d-a18e-f6908c068d63, status PAID e total 101.70.", "OK"),
        ],
        [1900, 5710, 1750],
    )

    heading(doc, "Kafka e Resiliencia", 1)
    paragraph(
        doc,
        "Os topicos Kafka encontrados confirmam o caminho principal e a infraestrutura de retry/DLT para processamento de pagamentos.",
    )
    add_matrix_table(
        doc,
        ["Topico", "Uso observado"],
        [
            ("payment-requests", "Entrada assincrona dos pedidos criados pelo order-service."),
            ("payment-results", "Saida de resultado consumida pelo order-service para atualizar o pedido."),
            ("payment-requests-retry-0", "Primeira etapa de retry configurada no processamento de pagamentos."),
            ("payment-requests-retry-1", "Segunda etapa de retry configurada no processamento de pagamentos."),
            ("payment-requests.DLT", "Destino final para mensagens nao processadas apos tentativas."),
            ("__consumer_offsets", "Topico interno Kafka para controle de offsets dos consumidores."),
        ],
        [3100, 6260],
    )

    heading(doc, "Observacoes", 1)
    add_matrix_table(
        doc,
        ["Item", "Observacao"],
        [
            ("Avisos Maven", "Foram observados warnings de Mockito/ByteBuddy com JDK 25 e avisos de APIs depreciadas em dependencias; nao bloquearam os testes."),
            ("Eureka em testes", "Com o ambiente Docker ativo, o gateway conseguiu consultar o Eureka durante o teste de contexto."),
            ("Estado final", "Os containers permaneceram em execucao para inspecao manual apos a validacao."),
        ],
        [2200, 7160],
    )

    heading(doc, "Conclusao", 1)
    paragraph(
        doc,
        "O projeto foi executado com sucesso localmente. A autenticacao centralizada no gateway, a descoberta via Eureka, a criacao de pedidos, o processamento assincrono de pagamentos por Kafka, os topicos de resiliencia e a atualizacao persistida do pedido foram confirmados em ambiente Docker Compose.",
    )

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
